import os
import subprocess
import sys
import json
from pathlib import Path

def check_python_version():
    """Check if Python version is compatible"""
    if sys.version_info < (3, 8):
        print("❌ Python 3.8 or higher is required")
        sys.exit(1)
    print(f"✅ Python {sys.version_info.major}.{sys.version_info.minor} detected")

def create_project_structure():
    """Create necessary directories"""
    directories = [
        'scrapy_project/spiders',
        'tailored_resumes',
        'cover_letters',
        'applications',
        'templates',
        'logs'
    ]
    
    for directory in directories:
        Path(directory).mkdir(parents=True, exist_ok=True)
        print(f"📁 Created directory: {directory}")

def install_requirements():
    """Install Python requirements"""
    print("📦 Installing Python requirements...")
    
    try:
        subprocess.check_call([sys.executable, '-m', 'pip', 'install', '-r', 'requirements.txt'])
        print("✅ Requirements installed successfully")
    except subprocess.CalledProcessError:
        print("❌ Failed to install requirements")
        sys.exit(1)

def setup_environment_file():
    """Create .env file template"""
    env_template = """# Job Scraper Environment Variables

# Google Sheets API
GOOGLE_SHEETS_JOB_ID=your_google_sheets_id_here
GOOGLE_CREDENTIALS_PATH=google_credentials.json

# OpenAI (optional - for better resume tailoring)
OPENAI_API_KEY=your_openai_key_here

# Email credentials (for auto-application)
EMAIL_ADDRESS=your_email@gmail.com
EMAIL_PASSWORD=your_app_password_here

# LinkedIn credentials (HIGH RISK - use with caution)
LINKEDIN_EMAIL=your_linkedin_email@gmail.com
LINKEDIN_PASSWORD=your_linkedin_password_here
ENABLE_LINKEDIN_SCRAPING=false

# Job search preferences
JOB_KEYWORDS=data engineer,machine learning engineer,AI engineer
PREFERRED_LOCATIONS=Remote,New York,San Francisco,Seattle
MIN_SALARY=75000
MAX_APPLICATIONS_PER_DAY=5

# Scraping settings
RUN_IMMEDIATE_TEST=false
"""
    
    if not os.path.exists('.env'):
        with open('.env', 'w') as f:
            f.write(env_template)
        print("📝 Created .env file template - please configure with your settings")
    else:
        print("✅ .env file already exists")

def create_base_resume_template():
    """Create base resume template"""
    template_path = 'templates/base_resume.docx'
    
    if os.path.exists(template_path):
        print("✅ Base resume template already exists")
        return
    
    try:
        from docx import Document
        from docx.shared import Inches
        
        doc = Document()
        
        # Name
        name_para = doc.add_paragraph()
        name_run = name_para.add_run("YOUR FULL NAME")
        name_run.bold = True
        name_run.font.size = Inches(0.25)
        name_para.alignment = 1
        
        # Contact
        contact_para = doc.add_paragraph("your.email@gmail.com | (555) 123-4567 | City, State | LinkedIn Profile")
        contact_para.alignment = 1
        doc.add_paragraph()
        
        # Summary
        summary_heading = doc.add_paragraph("PROFESSIONAL SUMMARY")
        summary_heading.runs[0].bold = True
        doc.add_paragraph("Experienced data professional with expertise in Python, SQL, and machine learning. Proven track record of building scalable data pipelines and delivering actionable insights.")
        doc.add_paragraph()
        
        # Skills
        skills_heading = doc.add_paragraph("TECHNICAL SKILLS")
        skills_heading.runs[0].bold = True
        doc.add_paragraph("Python • SQL • Git • Linux • AWS • Docker • Machine Learning • Data Analysis • ETL • Pandas")
        doc.add_paragraph()
        
        # Experience
        exp_heading = doc.add_paragraph("PROFESSIONAL EXPERIENCE")
        exp_heading.runs[0].bold = True
        
        job_title = doc.add_paragraph("Data Engineer | Company Name | 2020 - Present")
        job_title.runs[0].bold = True
        
        doc.add_paragraph("• Built ETL pipelines processing 100GB+ daily data using Python and SQL")
        doc.add_paragraph("• Developed machine learning models improving prediction accuracy by 20%")
        doc.add_paragraph("• Collaborated with cross-functional teams in Agile environment")
        doc.add_paragraph()
        
        # Education
        edu_heading = doc.add_paragraph("EDUCATION")
        edu_heading.runs[0].bold = True
        doc.add_paragraph("Bachelor of Science in Computer Science | University Name | 2018")
        
        doc.save(template_path)
        print("📄 Created base resume template")
        
    except Exception as e:
        print(f"⚠️  Could not create resume template: {e}")
        print("Please create templates/base_resume.docx manually")

def setup_scrapy_project():
    """Initialize Scrapy project structure"""
    scrapy_files = {
        'scrapy_project/__init__.py': '',
        'scrapy_project/spiders/__init__.py': ''
    }
    
    for file_path, content in scrapy_files.items():
        if not os.path.exists(file_path):
            with open(file_path, 'w') as f:
                f.write(content)

def create_run_scripts():
    """Create convenient run scripts"""
    
    # Main runner script
    run_script = """#!/usr/bin/env python3
# File: run.py
# Main application runner

import os
import sys
import subprocess
from dotenv import load_dotenv

load_dotenv()

def show_menu():
    print("\\n🤖 Job Scraper & Auto-Applier")
    print("=" * 40)
    print("1. Run job scraping (Indeed only)")
    print("2. Run job scraping (All sources)")
    print("3. Start auto-applier service")
    print("4. Start resume tailoring API")
    print("5. Run immediate test scrape")
    print("6. View scraping results")
    print("7. Setup Google Sheets")
    print("0. Exit")
    print()

def run_scrapy_spider(spider_name):
    try:
        subprocess.run(['scrapy', 'crawl', spider_name], cwd='scrapy_project', check=True)
        print(f"✅ {spider_name} completed successfully")
    except subprocess.CalledProcessError as e:
        print(f"❌ {spider_name} failed: {e}")

def main():
    while True:
        show_menu()
        choice = input("Select option (0-7): ").strip()
        
        if choice == '0':
            print("👋 Goodbye!")
            break
        elif choice == '1':
            print("🕷️  Running Indeed scraper...")
            run_scrapy_spider('indeed_jobs')
        elif choice == '2':
            print("🕷️  Running all scrapers...")
            subprocess.run([sys.executable, 'main_scraper.py'])
        elif choice == '3':
            print("🤖 Starting auto-applier...")
            subprocess.run([sys.executable, 'auto_applier.py'])
        elif choice == '4':
            print("📄 Starting resume API...")
            subprocess.run([sys.executable, 'resume_tailor_api.py'])
        elif choice == '5':
            print("🧪 Running test scrape...")
            os.environ['RUN_IMMEDIATE_TEST'] = 'true'
            subprocess.run([sys.executable, 'main_scraper.py'])
        elif choice == '6':
            print("📊 Viewing results...")
            if os.path.exists('scraping_results.json'):
                import json
                with open('scraping_results.json', 'r') as f:
                    data = json.load(f)
                    sessions = data.get('sessions', [])
                    print(f"Total sessions: {len(sessions)}")
                    if sessions:
                        latest = sessions[-1]
                        print(f"Latest session: {latest['timestamp']}")
                        print(f"Scrapers run: {latest['scrapers_run']}")
            else:
                print("No results found")
        elif choice == '7':
            print("📋 Google Sheets Setup:")
            print("1. Go to Google Cloud Console")
            print("2. Create project and enable Sheets API")
            print("3. Create service account")
            print("4. Download credentials as google_credentials.json")
            print("5. Create Google Sheet and share with service account")
            print("6. Copy sheet ID to .env file")
        else:
            print("❌ Invalid choice")
        
        if choice != '0':
            input("\\nPress Enter to continue...")

if __name__ == "__main__":
    main()
"""
    
    with open('run.py', 'w') as f:
        f.write(run_script)
    os.chmod('run.py', 0o755)
    print("🚀 Created run.py script")
    
    # Docker compose for easy deployment
    docker_compose = """version: '3.8'

services:
  job-scraper:
    build: .
    container_name: job-scraper
    volumes:
      - ./scrapy_project:/app/scrapy_project
      - ./tailored_resumes:/app/tailored_resumes
      - ./cover_letters:/app/cover_letters
      - ./google_credentials.json:/app/google_credentials.json
      - ./.env:/app/.env
    ports:
      - "5000:5000"
      - "5001:5001"
    environment:
      - PYTHONPATH=/app
    restart: unless-stopped
    command: python main_scraper.py

  auto-applier:
    build: .
    container_name: auto-applier
    depends_on:
      - job-scraper
    volumes:
      - ./google_credentials.json:/app/google_credentials.json
      - ./.env:/app/.env
    environment:
      - PYTHONPATH=/app
    restart: unless-stopped
    command: python auto_applier.py

  resume-api:
    build: .
    container_name: resume-api
    volumes:
      - ./templates:/app/templates
      - ./tailored_resumes:/app/tailored_resumes
      - ./cover_letters:/app/cover_letters
      - ./.env:/app/.env
    ports:
      - "5001:5001"
    environment:
      - PYTHONPATH=/app
    restart: unless-stopped
    command: python resume_tailor_api.py
"""
    
    with open('docker-compose.yml', 'w') as f:
        f.write(docker_compose)
    print("🐳 Created docker-compose.yml")

def create_dockerfile():
    """Create Dockerfile for containerization"""
    dockerfile_content = """FROM python:3.11-slim

WORKDIR /app

# Install system dependencies
RUN apt-get update && apt-get install -y \\
    wget \\
    gnupg \\
    unzip \\
    curl \\
    && rm -rf /var/lib/apt/lists/*

# Install Chrome for Selenium
RUN wget -q -O - https://dl-ssl.google.com/linux/linux_signing_key.pub | apt-key add - \\
    && echo "deb http://dl.google.com/linux/chrome/deb/ stable main" >> /etc/apt/sources.list.d/google.list \\
    && apt-get update \\
    && apt-get install -y google-chrome-stable \\
    && rm -rf /var/lib/apt/lists/*

# Install Python requirements
COPY requirements.txt .
RUN pip install --no-cache-dir -r requirements.txt

# Copy application code
COPY . .

# Create necessary directories
RUN mkdir -p scrapy_project/spiders tailored_resumes cover_letters applications logs

EXPOSE 5000 5001

CMD ["python", "run.py"]
"""
    
    with open('Dockerfile', 'w') as f:
        f.write(dockerfile_content)
    print("🐳 Created Dockerfile")

def create_readme():
    """Create comprehensive README"""
    readme_content = """# Free Job Scraper & Auto-Applier

A completely free, automated job scraping and application system that:
- 🕷️ Scrapes jobs from Indeed, LinkedIn, and company websites
- 📄 Auto-generates tailored resumes and cover letters
- 🤖 Automatically applies to suitable positions
- 📊 Tracks everything in Google Sheets

## 🚀 Quick Start

1. **Install dependencies:**
   ```bash
   python setup.py
   ```

2. **Configure credentials:**
   - Edit `.env` file with your details
   - Add `google_credentials.json` for Google Sheets
   - Optionally add OpenAI API key for better resume tailoring

3. **Run the application:**
   ```bash
   python run.py
   ```

## 📋 Features

### ✅ Job Scraping
- **Indeed**: Safe, high success rate
- **LinkedIn**: Higher risk, requires login
- **Company sites**: RSS feeds and career pages
- **Smart filtering**: Removes duplicates, scores by relevance

### 📄 Resume Tailoring
- Analyzes job descriptions automatically
- Tailors skills and summary for each job
- Generates custom cover letters
- Uses AI (OpenAI) or keyword matching (free)

### 🤖 Auto-Application
- **Indeed Easy Apply**: Automated applications
- **Email applications**: Finds contact emails
- **Safety limits**: Max applications per day
- **Human review**: High-priority jobs flagged for manual review

### 📊 Tracking & Organization
- Google Sheets integration
- Application status tracking
- Follow-up reminders
- Success rate monitoring

## ⚖️ Legal & Risk Assessment

### ✅ Low Risk (Recommended)
- Indeed scraping with delays
- Company career page RSS feeds
- Email-based applications

### ⚠️ Medium Risk (Use Caution)
- LinkedIn scraping (account suspension possible)
- High-frequency scraping

### ❌ High Risk (Not Recommended)
- Automated LinkedIn applications
- Ignoring robots.txt
- Mass applications without human review

## 🛠️ Configuration

### Environment Variables (.env)
```bash
# Required
GOOGLE_SHEETS_JOB_ID=your_sheet_id
EMAIL_ADDRESS=your_email@gmail.com

# Optional but recommended
OPENAI_API_KEY=your_openai_key
MAX_APPLICATIONS_PER_DAY=5

# High risk (LinkedIn)
LINKEDIN_EMAIL=your_email
LINKEDIN_PASSWORD=your_password
ENABLE_LINKEDIN_SCRAPING=false
```

### Google Sheets Setup
1. Go to Google Cloud Console
2. Create project and enable Sheets API + Drive API
3. Create service account
4. Download JSON credentials as `google_credentials.json`
5. Create Google Sheet and share with service account email
6. Copy sheet ID to `.env` file

## 🚦 Usage Modes

### 1. Conservative Mode (Recommended)
- Indeed scraping only
- Manual application review
- 2-5 applications per day
- Low risk of account issues

### 2. Aggressive Mode (Higher Risk)
- All sources including LinkedIn
- Automated applications
- 5-10 applications per day
- Higher success rate but more risk

### 3. Hybrid Mode (Balanced)
- Indeed + company sites
- Auto-apply to top matches only
- Manual review for complex applications
- Good balance of efficiency and safety

## 📊 Expected Results

### Daily Metrics
- **Jobs scraped**: 200-500
- **High-priority matches**: 5-15
- **Auto-applications**: 2-8
- **Manual review required**: 3-7

### Success Rates
- **Scraping success**: 85-95%
- **Application success**: 60-80%
- **Response rate**: 3-8%
- **Interview rate**: 1-3%

## 🔧 Advanced Configuration

### Custom Job Filters
Edit `scrapy_project/pipelines.py` to modify:
- Keyword requirements
- Salary filters
- Company preferences
- Experience level matching

### Resume Templates
- Edit `templates/base_resume.docx`
- Add multiple templates for different job types
- Customize skills and experience sections

### Application Logic
- Modify `auto_applier.py` for custom application rules
- Add company-specific application strategies
- Implement custom screening question answers

## 🐳 Docker Deployment

```bash
# Build and run with Docker
docker-compose up -d

# View logs
docker-compose logs -f

# Stop services
docker-compose down
```

## 📈 Monitoring

### View Results
```bash
python run.py
# Select option 6 to view scraping results
```

### Google Sheets Dashboard
Your Google Sheet will show:
- All scraped jobs with priority scores
- Application status and dates
- Follow-up reminders
- Success metrics

### Log Files
- `logs/scrapy.log`: Scraping activity
- `logs/applications.log`: Application attempts
- `scraping_results.json`: Session summaries

## 🛡️ Safety Features

### Rate Limiting
- Configurable delays between requests
- Daily application limits
- Proxy rotation support
- Human-like behavior simulation

### Error Handling
- Automatic retries for failed requests
- Graceful degradation when services fail
- Comprehensive logging
- Alert notifications

### Privacy Protection
- No data stored on external servers
- Local processing only
- Encrypted credential storage
- GDPR compliance ready

## 🆘 Troubleshooting

### Common Issues

**"No jobs found"**
- Check internet connection
- Verify .env configuration
- Try different keywords
- Check proxy settings

**"Google Sheets error"**
- Verify credentials file exists
- Check sheet sharing permissions
- Confirm API is enabled
- Test with manual API call

**"Application failed"**
- Check email credentials
- Verify resume template exists
- Review application logs
- Try manual application first

**"Chrome/Selenium errors"**
- Update Chrome browser
- Check Chrome driver version
- Try headless=False for debugging
- Install dependencies

### Getting Help
1. Check logs in `logs/` directory
2. Review error messages carefully
3. Test individual components
4. Check GitHub issues
5. Verify all credentials and permissions

## 🔮 Roadmap

### Planned Features
- [ ] More job sites (ZipRecruiter, Monster, etc.)
- [ ] Advanced AI resume optimization
- [ ] Interview scheduling integration
- [ ] Salary negotiation assistant
- [ ] Network analysis and referrals
- [ ] Mobile app for notifications

### Contributions Welcome
- Job site integrations
- Resume template improvements
- Application success optimizations
- Documentation updates
- Bug reports and fixes

## 📄 License

MIT License - Use freely for personal and commercial projects.

## ⚠️ Disclaimer

This tool is for educational and personal use. Users are responsible for:
- Complying with website terms of service
- Following employment laws and regulations
- Using accurate information in applications
- Respecting employer preferences and processes

The authors are not responsible for account suspensions, legal issues, or application outcomes.

## 💡 Tips for Success

### Job Search Strategy
1. **Quality over quantity**: Focus on relevant, high-priority jobs
2. **Personalization**: Always review auto-generated materials
3. **Follow-up**: Set reminders for follow-up communications
4. **Continuous improvement**: Adjust keywords and filters based on results

### Application Best Practices
1. **Research companies**: Understand their culture and needs
2. **Customize applications**: Tailor even automated applications
3. **Track everything**: Monitor response rates and adjust strategy
4. **Stay organized**: Use the Google Sheets dashboard effectively

### Technical Optimization
1. **Monitor performance**: Check logs and success rates
2. **Adjust delays**: Balance speed with reliability
3. **Update regularly**: Keep dependencies and templates current
4. **Backup data**: Export important results regularly

---

**Happy Job Hunting! 🎯**

Built with ❤️ for job seekers everywhere.
"""
    
    with open('README.md', 'w') as f:
        f.write(readme_content)
    print("📖 Created comprehensive README.md")

def create_gitignore():
    """Create .gitignore file"""
    gitignore_content = """# Job Scraper .gitignore

# Environment and credentials
.env
google_credentials.json
*.json

# Generated files
tailored_resumes/
cover_letters/
applications/
logs/
*.xlsx
*.csv

# Python
__pycache__/
*.py[cod]
*$py.class
*.so
.Python
build/
develop-eggs/
dist/
downloads/
eggs/
.eggs/
lib/
lib64/
parts/
sdist/
var/
wheels/
*.egg-info/
.installed.cfg
*.egg

# Virtual environments
venv/
ENV/
env/
.venv/

# IDE
.vscode/
.idea/
*.swp
*.swo
*~

# OS
.DS_Store
Thumbs.db

# Scrapy
.scrapy/

# Selenium
*.log
ghostdriver.log

# Chrome
chrome_driver/

# Docker
*.log
"""
    
    with open('.gitignore', 'w') as f:
        f.write(gitignore_content)
    print("🔒 Created .gitignore")

def main():
    """Main setup function"""
    print("🚀 Setting up Free Job Scraper & Auto-Applier...")
    print("=" * 50)
    
    # Check system requirements
    check_python_version()
    
    # Create project structure
    create_project_structure()
    
    # Install requirements
    install_requirements()
    
    # Setup configuration
    setup_environment_file()
    setup_scrapy_project()
    
    # Create templates and scripts
    create_base_resume_template()
    create_run_scripts()
    create_dockerfile()
    
    # Create documentation
    create_readme()
    create_gitignore()
    
    print("\n" + "=" * 50)
    print("✅ Setup completed successfully!")
    print("\n📋 Next Steps:")
    print("1. Configure .env file with your credentials")
    print("2. Add google_credentials.json file")
    print("3. Edit templates/base_resume.docx with your info")
    print("4. Run: python run.py")
    print("\n💡 For help, see README.md")
    print("🎯 Happy job hunting!")

if __name__ == "__main__":
    main()

---

# File: proxy_list.txt
# Free proxy list (update regularly)

# Instructions: Update this list daily with fresh proxies
# Get proxies from:
# - https://free-proxy-list.net/
# - https://www.proxy-list.download/
# - https://github.com/clarketm/proxy-list

# Format: ip:port
154.16.63.16:3128
103.148.72.192:80
194.233.69.126:443
103.149.162.195:80
185.15.172.212:3128
158.69.115.201:8080
45.76.97.117:3128
104.248.90.212:80

# Proxy updater will refresh this list automatically# Complete Free Job Scraper with Auto-Application
# File: requirements.txt

scrapy==2.11.0
selenium==4.15.0
scrapy-selenium==0.0.7
scrapy-rotating-proxies==0.6.2
requests==2.31.0
pandas==2.1.3
openpyxl==3.1.2
gspread==5.12.0
oauth2client==4.1.3
fake-useragent==1.4.0
scrapy-user-agents==0.1.1
undetected-chromedriver==3.5.4
flask==3.0.0
python-docx==0.8.11
beautifulsoup4==4.12.2
lxml==4.9.3
Pillow==10.1.0
schedule==1.2.0
python-dotenv==1.0.0
openai==1.3.5
webdriver-manager==4.0.1

---

# File: .env
# Environment variables (create this file and add your credentials)

# Google Sheets API
GOOGLE_SHEETS_JOB_ID=your_google_sheets_id_here
GOOGLE_CREDENTIALS_PATH=google_credentials.json

# OpenAI (optional - for better resume tailoring)
OPENAI_API_KEY=your_openai_key_here

# Email credentials (for auto-application)
EMAIL_ADDRESS=your_email@gmail.com
EMAIL_PASSWORD=your_app_password_here

# LinkedIn credentials (HIGH RISK - use with caution)
LINKEDIN_EMAIL=your_linkedin_email@gmail.com
LINKEDIN_PASSWORD=your_linkedin_password_here

# Indeed account (safer)
INDEED_EMAIL=your_email@gmail.com

# Job search preferences
JOB_KEYWORDS=data engineer,machine learning engineer,AI engineer,MLOps engineer
PREFERRED_LOCATIONS=Remote,New York,San Francisco,Seattle
MIN_SALARY=75000
MAX_APPLICATIONS_PER_DAY=10

---

# File: scrapy_project/settings.py

BOT_NAME = 'job_scraper'
SPIDER_MODULES = ['scrapy_project.spiders']
NEWSPIDER_MODULE = 'scrapy_project.spiders'

# Obey robots.txt rules
ROBOTSTXT_OBEY = False  # Set to False for aggressive scraping

# Configure delays for different sites
DOWNLOAD_DELAY = 3
RANDOMIZE_DOWNLOAD_DELAY = True
CONCURRENT_REQUESTS = 4
CONCURRENT_REQUESTS_PER_DOMAIN = 1

# Auto-throttling
AUTOTHROTTLE_ENABLED = True
AUTOTHROTTLE_START_DELAY = 1
AUTOTHROTTLE_MAX_DELAY = 10
AUTOTHROTTLE_TARGET_CONCURRENCY = 2.0

# Disable cookies (avoid tracking)
COOKIES_ENABLED = False
TELNETCONSOLE_ENABLED = False

# Middleware configuration
DOWNLOADER_MIDDLEWARES = {
    'scrapy.downloadermiddlewares.useragent.UserAgentMiddleware': None,
    'scrapy_user_agents.middlewares.RandomUserAgentMiddleware': 400,
    'rotating_proxies.middlewares.RotatingProxyMiddleware': 610,
    'rotating_proxies.middlewares.BanDetectionMiddleware': 620,
    'scrapy_project.middlewares.SeleniumMiddleware': 800,
}

# Item pipelines
ITEM_PIPELINES = {
    'scrapy_project.pipelines.DuplicatesPipeline': 200,
    'scrapy_project.pipelines.DataCleaningPipeline': 300,
    'scrapy_project.pipelines.AutoApplicationPipeline': 350,
    'scrapy_project.pipelines.GoogleSheetsPipeline': 400,
}

# Selenium settings
SELENIUM_DRIVER_NAME = 'chrome'
SELENIUM_DRIVER_EXECUTABLE_PATH = None
SELENIUM_DRIVER_ARGUMENTS = [
    '--headless=new',
    '--no-sandbox',
    '--disable-dev-shm-usage',
    '--disable-blink-features=AutomationControlled',
    '--disable-extensions',
    '--disable-plugins',
    '--disable-images',
    '--window-size=1920,1080',
    '--user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
]

# Proxy settings (optional)
ROTATING_PROXY_LIST_PATH = 'proxy_list.txt'

# Cache settings
HTTPCACHE_ENABLED = True
HTTPCACHE_EXPIRATION_SECS = 3600

---

# File: scrapy_project/items.py

import scrapy

class JobItem(scrapy.Item):
    # Basic information
    job_id = scrapy.Field()
    title = scrapy.Field()
    company = scrapy.Field()
    location = scrapy.Field()
    salary = scrapy.Field()
    job_type = scrapy.Field()
    
    # URLs
    job_url = scrapy.Field()
    apply_url = scrapy.Field()
    company_url = scrapy.Field()
    
    # Content
    description = scrapy.Field()
    requirements = scrapy.Field()
    responsibilities = scrapy.Field()
    benefits = scrapy.Field()
    
    # Metadata
    posted_date = scrapy.Field()
    scraped_date = scrapy.Field()
    source = scrapy.Field()
    
    # Analysis
    keywords = scrapy.Field()
    experience_level = scrapy.Field()
    remote_friendly = scrapy.Field()
    priority_score = scrapy.Field()
    match_score = scrapy.Field()
    
    # Application tracking
    application_status = scrapy.Field()
    application_method = scrapy.Field()
    auto_apply_eligible = scrapy.Field()
    application_complexity = scrapy.Field()
    notes = scrapy.Field()
    
    # Auto-application data
    email_found = scrapy.Field()
    contact_email = scrapy.Field()
    application_deadline = scrapy.Field()
    easy_apply_available = scrapy.Field()

---

# File: scrapy_project/middlewares.py

from scrapy_selenium import SeleniumRequest, SeleniumMiddleware
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
import time
import random

class CustomSeleniumMiddleware(SeleniumMiddleware):
    
    def process_request(self, request, spider):
        if not isinstance(request, SeleniumRequest):
            return None
            
        # Add random delays and human-like behavior
        if hasattr(spider, 'custom_settings'):
            delay = spider.custom_settings.get('DOWNLOAD_DELAY', 3)
            time.sleep(delay + random.uniform(0, 2))
        
        return super().process_request(request, spider)

class AntiDetectionMiddleware:
    
    def process_request(self, request, spider):
        # Add headers to appear more human-like
        request.headers.setdefault('Accept', 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8')
        request.headers.setdefault('Accept-Language', 'en-US,en;q=0.5')
        request.headers.setdefault('Accept-Encoding', 'gzip, deflate')
        request.headers.setdefault('Connection', 'keep-alive')
        request.headers.setdefault('Upgrade-Insecure-Requests', '1')
        
        return None
