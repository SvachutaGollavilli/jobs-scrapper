from flask import Flask, request, jsonify
import os
import openai
from docx import Document
from docx.shared import Inches
import tempfile
from datetime import datetime
import re
import requests
from bs4 import BeautifulSoup
from dotenv import load_dotenv

load_dotenv()

app = Flask(__name__)

class ResumeTailorer:
    def __init__(self):
        self.openai_key = os.getenv('OPENAI_API_KEY')
        if self.openai_key:
            openai.api_key = self.openai_key
    
    def scrape_job_description(self, job_url):
        """Scrape job description from URL"""
        try:
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
            }
            response = requests.get(job_url, headers=headers, timeout=10)
            soup = BeautifulSoup(response.content, 'html.parser')
            
            # Try multiple selectors for job descriptions
            selectors = [
                '[data-testid="jobsearch-JobComponent-description"]',
                '.jobs-description__content',
                '#jobDescriptionText',
                '.job-description',
                '.jobsearch-jobDescriptionText'
            ]
            
            for selector in selectors:
                element = soup.select_one(selector)
                if element:
                    return element.get_text(strip=True)
            
            # Fallback: get body text
            return soup.get_text(strip=True)[:5000]
            
        except Exception as e:
            print(f"Error scraping job description: {e}")
            return ""
    
    def analyze_job_with_ai(self, job_description):
        """Analyze job using OpenAI"""
        if not self.openai_key:
            return self.analyze_job_with_keywords(job_description)
        
        try:
            response = openai.ChatCompletion.create(
                model="gpt-3.5-turbo",
                messages=[
                    {
                        "role": "system",
                        "content": "You are a job analysis expert. Extract key requirements from job descriptions and return structured JSON."
                    },
                    {
                        "role": "user",
                        "content": f"""
                        Analyze this job description and extract:
                        
                        Job Description:
                        {job_description[:3000]}
                        
                        Return JSON with:
                        {{
                            "technical_skills": ["skill1", "skill2"],
                            "soft_skills": ["skill1", "skill2"],
                            "experience_level": "junior/mid/senior",
                            "key_requirements": ["req1", "req2"],
                            "nice_to_have": ["skill1", "skill2"],
                            "keywords_to_emphasize": ["keyword1", "keyword2"],
                            "company_culture": "description"
                        }}
                        """
                    }
                ],
                temperature=0.1,
                max_tokens=1000
            )
            
            import json
            return json.loads(response.choices[0].message.content)
            
        except Exception as e:
            print(f"OpenAI analysis failed: {e}, falling back to keyword analysis")
            return self.analyze_job_with_keywords(job_description)
    
    def analyze_job_with_keywords(self, job_description):
        """Analyze job using keyword matching (free alternative)"""
        text = job_description.lower()
        
        # Technical skills mapping
        tech_skills_map = {
            'Python': ['python', 'py', 'pandas', 'numpy', 'scipy'],
            'SQL': ['sql', 'mysql', 'postgresql', 'database', 'queries'],
            'AWS': ['aws', 'amazon web services', 's3', 'ec2', 'lambda'],
            'Docker': ['docker', 'container', 'containerization'],
            'Kubernetes': ['kubernetes', 'k8s', 'orchestration'],
            'Machine Learning': ['machine learning', 'ml', 'ai', 'artificial intelligence'],
            'TensorFlow': ['tensorflow', 'tf', 'keras'],
            'PyTorch': ['pytorch', 'torch'],
            'Apache Spark': ['spark', 'apache spark', 'pyspark'],
            'Airflow': ['airflow', 'workflow', 'pipeline'],
            'Git': ['git', 'github', 'version control', 'gitlab'],
            'Linux': ['linux', 'unix', 'bash', 'shell'],
            'Java': ['java', 'jvm', 'spring'],
            'Scala': ['scala'],
            'R': [' r ', 'r programming', 'rstudio'],
            'Tableau': ['tableau', 'data visualization'],
            'Power BI': ['power bi', 'powerbi'],
            'Excel': ['excel', 'spreadsheet'],
            'NoSQL': ['nosql', 'mongodb', 'cassandra'],
            'Elasticsearch': ['elasticsearch', 'elk', 'kibana']
        }
        
        # Soft skills mapping
        soft_skills_map = {
            'Communication': ['communication', 'collaborate', 'teamwork'],
            'Leadership': ['leadership', 'lead', 'mentor', 'manage'],
            'Problem Solving': ['problem solving', 'analytical', 'critical thinking'],
            'Project Management': ['project management', 'agile', 'scrum']
        }
        
        # Find technical skills
        found_tech_skills = []
        for skill, keywords in tech_skills_map.items():
            if any(keyword in text for keyword in keywords):
                found_tech_skills.append(skill)
        
        # Find soft skills
        found_soft_skills = []
        for skill, keywords in soft_skills_map.items():
            if any(keyword in text for keyword in keywords):
                found_soft_skills.append(skill)
        
        # Determine experience level
        experience_level = 'mid'
        if any(term in text for term in ['senior', 'lead', 'principal', '5+ years', '6+ years']):
            experience_level = 'senior'
        elif any(term in text for term in ['junior', 'entry', 'graduate', '0-2 years']):
            experience_level = 'junior'
        
        # Extract key requirements
        sentences = job_description.split('.')
        requirements = []
        for sentence in sentences:
            if any(phrase in sentence.lower() for phrase in ['required', 'must have', 'essential']):
                requirements.append(sentence.strip())
                if len(requirements) >= 5:
                    break
        
        return {
            'technical_skills': found_tech_skills[:10],
            'soft_skills': found_soft_skills[:5],
            'experience_level': experience_level,
            'key_requirements': requirements,
            'nice_to_have': [],
            'keywords_to_emphasize': found_tech_skills[:8],
            'company_culture': 'Data-driven, collaborative environment'
        }
    
    def tailor_resume_content(self, base_resume_path, job_analysis, job_info):
        """Tailor resume based on job analysis"""
        try:
            # Load base resume
            doc = Document(base_resume_path)
            
            # Extract current content
            resume_text = '\n'.join([p.text for p in doc.paragraphs])
            
            # Create new tailored resume
            new_doc = Document()
            
            # Add tailored content
            self.add_tailored_header(new_doc, job_info)
            self.add_tailored_summary(new_doc, job_analysis, resume_text)
            self.add_tailored_skills(new_doc, job_analysis, resume_text)
            self.add_experience_section(new_doc, resume_text, job_analysis)
            self.add_education_section(new_doc, resume_text)
            
            # Save tailored resume
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            company = job_info.get('company', 'Company').replace(' ', '_')
            output_path = f"tailored_resumes/{company}_resume_{timestamp}.docx"
            
            os.makedirs('tailored_resumes', exist_ok=True)
            new_doc.save(output_path)
            
            return output_path
            
        except Exception as e:
            print(f"Error tailoring resume: {e}")
            return None
    
    def add_tailored_header(self, doc, job_info):
        """Add header section"""
        # Name
        name_para = doc.add_paragraph()
        name_run = name_para.add_run("YOUR NAME")
        name_run.bold = True
        name_run.font.size = Inches(0.3)
        name_para.alignment = 1  # Center
        
        # Contact info
        contact_para = doc.add_paragraph("your.email@gmail.com | (555) 123-4567 | LinkedIn Profile")
        contact_para.alignment = 1
        doc.add_paragraph()
    
    def add_tailored_summary(self, doc, job_analysis, resume_text):
        """Add tailored professional summary"""
        heading = doc.add_paragraph("PROFESSIONAL SUMMARY")
        heading.runs[0].bold = True
        
        # Create tailored summary
        keywords = job_analysis.get('keywords_to_emphasize', [])[:5]
        experience_level = job_analysis.get('experience_level', 'experienced')
        
        summary = f"Results-driven {experience_level} professional with expertise in {', '.join(keywords[:3])}. "
        summary += f"Proven track record in data engineering and machine learning with hands-on experience in {', '.join(keywords[3:5])}. "
        summary += "Strong analytical and problem-solving skills with a passion for turning data into actionable insights."
        
        doc.add_paragraph(summary)
        doc.add_paragraph()
    
    def add_tailored_skills(self, doc, job_analysis, resume_text):
        """Add tailored technical skills section"""
        heading = doc.add_paragraph("TECHNICAL SKILLS")
        heading.runs[0].bold = True
        
        # Combine job requirements with existing skills
        job_skills = job_analysis.get('technical_skills', [])
        
        # Base skills that most data engineers have
        base_skills = ['Python', 'SQL', 'Git', 'Linux', 'Excel']
        
        # Combine and deduplicate
        all_skills = list(set(job_skills + base_skills))
        
        skills_text = ' • '.join(all_skills)
        doc.add_paragraph(skills_text)
        doc.add_paragraph()
    
    def add_experience_section(self, doc, resume_text, job_analysis):
        """Add professional experience section"""
        heading = doc.add_paragraph("PROFESSIONAL EXPERIENCE")
        heading.runs[0].bold = True
        
        # This would need to parse existing experience from resume_text
        # For now, add placeholder
        exp_para = doc.add_paragraph("Data Engineer | Previous Company | 2020 - Present")
        exp_para.runs[0].bold = True
        
        # Add bullet points emphasizing relevant skills
        keywords = job_analysis.get('keywords_to_emphasize', [])
        
        bullets = [
            f"Developed ETL pipelines using {keywords[0] if keywords else 'Python'} to process large datasets",
            f"Built machine learning models with {keywords[1] if len(keywords) > 1 else 'TensorFlow'} improving accuracy by 15%",
            f"Collaborated with cross-functional teams using {keywords[2] if len(keywords) > 2 else 'Agile'} methodologies",
        ]
        
        for bullet in bullets:
            doc.add_paragraph(f"• {bullet}")
        
        doc.add_paragraph()
    
    def add_education_section(self, doc, resume_text):
        """Add education section"""
        heading = doc.add_paragraph("EDUCATION")
        heading.runs[0].bold = True
        
        doc.add_paragraph("Bachelor of Science in Computer Science | University Name | 2018")
        doc.add_paragraph()
    
    def generate_cover_letter(self, job_analysis, job_info):
        """Generate tailored cover letter"""
        company = job_info.get('company', '[Company Name]')
        title = job_info.get('title', '[Position Title]')
        
        skills = ', '.join(job_analysis.get('technical_skills', [])[:3])
        
        cover_letter = f"""Dear Hiring Manager,

I am writing to express my strong interest in the {title} position at {company}. With my background in data engineering and machine learning, I am excited about the opportunity to contribute to your team's success.

Key qualifications that make me a strong fit:

• Technical expertise in {skills}, directly aligning with your requirements
• Proven experience in building scalable data pipelines and ETL processes
• Strong analytical skills and experience with machine learning model development
• Collaborative approach and excellent communication skills

I am particularly drawn to {company} because of your reputation for innovation and data-driven decision making. I believe my technical skills and passion for solving complex problems would make me a valuable addition to your team.

I would welcome the opportunity to discuss how my experience and enthusiasm can contribute to your continued success. Thank you for your consideration.

Best regards,
[Your Name]
[Your Phone Number]
[Your Email]

---
Position applied for: {title}
Application date: {datetime.now().strftime("%Y-%m-%d")}
"""
        
        # Save cover letter
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        company_clean = company.replace(' ', '_').replace('[', '').replace(']', '')
        output_path = f"cover_letters/{company_clean}_cover_letter_{timestamp}.txt"
        
        os.makedirs('cover_letters', exist_ok=True)
        with open(output_path, 'w') as f:
            f.write(cover_letter)
        
        return {
            'content': cover_letter,
            'file_path': output_path
        }

# Flask API endpoints
tailor_service = ResumeTailorer()

@app.route('/tailor-resume', methods=['POST'])
def tailor_resume():
    """Tailor resume for specific job"""
    data = request.get_json()
    
    job_url = data.get('job_url')
    job_description = data.get('job_description', '')
    company = data.get('company', '')
    title = data.get('title', '')
    base_resume_path = data.get('base_resume_path', 'templates/base_resume.docx')
    
    if not os.path.exists(base_resume_path):
        return jsonify({'error': 'Base resume template not found'}), 400
    
    try:
        # Get job description if not provided
        if not job_description and job_url:
            job_description = tailor_service.scrape_job_description(job_url)
        
        if not job_description:
            return jsonify({'error': 'Job description could not be obtained'}), 400
        
        # Analyze job requirements
        job_analysis = tailor_service.analyze_job_with_ai(job_description)
        
        # Tailor resume
        job_info = {'company': company, 'title': title, 'url': job_url}
        resume_path = tailor_service.tailor_resume_content(base_resume_path, job_analysis, job_info)
        
        if not resume_path:
            return jsonify({'error': 'Failed to create tailored resume'}), 500
        
        # Generate cover letter
        cover_letter = tailor_service.generate_cover_letter(job_analysis, job_info)
        
        return jsonify({
            'success': True,
            'resume_path': resume_path,
            'cover_letter_path': cover_letter['file_path'],
            'cover_letter_content': cover_letter['content'],
            'job_analysis': job_analysis,
            'created_at': datetime.now().isoformat()
        })
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/batch-tailor', methods=['POST'])
def batch_tailor():
    """Batch tailor resumes for multiple jobs"""
    data = request.get_json()
    jobs = data.get('jobs', [])
    
    if not jobs:
        return jsonify({'error': 'No jobs provided'}), 400
    
    results = []
    
    for job in jobs[:5]:  # Limit to 5 jobs per batch
        try:
            result = tailor_resume()
            results.append({
                'job': job,
                'result': result.get_json()
            })
        except Exception as e:
            results.append({
                'job': job,
                'error': str(e)
            })
    
    return jsonify({
        'batch_results': results,
        'processed': len(results),
        'timestamp': datetime.now().isoformat()
    })

@app.route('/health', methods=['GET'])
def health():
    """Health check endpoint"""
    return jsonify({
        'status': 'healthy',
        'service': 'Resume Tailoring API',
        'timestamp': datetime.now().isoformat(),
        'openai_configured': bool(tailor_service.openai_key)
    })

if __name__ == '__main__':
    app.run(debug=True, port=5001)
